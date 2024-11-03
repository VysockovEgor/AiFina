from docx import Document
import re

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.enums import TA_JUSTIFY
from io import BytesIO
from reportlab.lib.fonts import addMapping

promt={}

class text_convert:
    def __init__(self, id, chatbot,promt='') -> None:
            self.id, self.chatbot = id, chatbot
            self.paragraphs = None  
            self.promt=promt

    def set_promt(self, promt): self.promt = promt
    def message_text(self, message):
        result=''
        
        if len(message)>7000:
            for i in message.split('\n'):
                if i!='':
                    result+=self.convert(i.strip())
        else: result=self.convert(message)
        return result

    def main(self, path, pindex, mes=False):

        if self.paragraphs is None:  
            if not mes:

                if path.split('.')[-1] == 'txt':
                    with open(path, 'r', encoding='utf-8') as file:
                        self.paragraphs = [i for i in file.read().split('\n') if i.strip()!='']
                        b_par = []
                        g_par = ''
                        for i in self.paragraphs:
                            if len(g_par) + len(i) < 1000:
                                g_par += i
                            else:
                                if g_par!='':
                                    b_par.append(g_par)
                                    g_par = i
                                else:b_par.append(i)
                        if g_par!='':
                            b_par.append(g_par)
                        self.paragraphs = b_par

                elif path.split('.')[-1] == 'docx':
                    doc = Document(path)

                    self.paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != '']
                    b_par = []
                    g_par = ''
                    for i in self.paragraphs:
                        if len(g_par) + len(i) < 1000:
                            g_par += i
                        else:
                            if g_par != '':
                                b_par.append(g_par)
                                g_par = i
                            else:
                                b_par.append(i)
                    if g_par != '':
                        b_par.append(g_par)
                    self.paragraphs = b_par
            else:
                self.paragraphs = [i for i in path.split('\n') if i.strip()!='']
                b_par = []
                g_par = ''
                for i in self.paragraphs:
                    if len(g_par) + len(i) < 1000:
                        g_par += i
                    else:
                        if g_par != '':
                            b_par.append(g_par)
                            g_par = i
                        else:
                            b_par.append(i)
                if g_par != '':
                    b_par.append(g_par)
                self.paragraphs = b_par


        mes = ''
        result = ''
        for sentence in re.split(r'(?<=[.!?])\s+', self.paragraphs[pindex]):
            if len(mes) + len(sentence) < 500:
                mes += sentence
            else:
                if mes!='':
                    result += self.convert(mes)
                    mes = sentence
                else:
                    result += self.convert(sentence)
        if mes !='':
            result += self.convert(mes)

        return result, len(self.paragraphs)

    def convert(self,t):
        return self.chatbot.process_message_with_manual_cancel(self.id,t+ f'\n<|promt|>Ты - машина, которая умеет только сильно сокращать текст, сохранив его основные идеи, тон и стиль, а также смысловую связь с предыдущей частью текста. {self.promt}, выдай только сокращенный текст, не добавляя ничего<|endofpromt|>')+'\n'*2


def text_to_pdf(text):
    pdf_stream = BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=20 * mm,
                            bottomMargin=20 * mm)

    pdfmetrics.registerFont(TTFont('Helvetica', 'Helvetica.ttf'))
    addMapping('Helvetica', 0, 0, 'Helvetica')

    styles = getSampleStyleSheet()
    styles['Normal'].fontName = 'Helvetica'
    styles['Normal'].fontSize = 12
    styles['Normal'].leading = 14
    styles['Normal'].alignment = TA_JUSTIFY

    story = []

    text_with_breaks = text.replace('\n', '<br/>')
    paragraph = Paragraph(text_with_breaks, styles['Normal'])
    story.append(paragraph)

    doc.build(story)
    pdf_stream.seek(0)
    return pdf_stream

'''chatbot = Chatbot()               

text, plen = text_convert(1,chatbot).main('1.docx', 0)
print(text)
for i in range(1,plen):
    text, plen = text_convert(1,chatbot).main('1.docx', i)
    print('='*50)
    print(text)'''